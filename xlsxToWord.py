from docx import Document
import win32com.client as win32
import openpyxl
#https://openpyxl.readthedocs.io/en/stable/api/openpyxl.worksheet.worksheet.html#openpyxl.worksheet.worksheet.Worksheet
#https://python-docx.readthedocs.io/en/latest/user/styles-using.html

#Open the word document
doc = Document('./InputChecker.docx') #replace with any path or make it an input later

#find a way to write from the xlsx file to the word document

#perhaps iterate through the xlsx file first

  
# load excel with its path 
wrkbk = openpyxl.load_workbook(r"./Employee Sample Data.xlsx") 
  
sh = wrkbk.active #active workbook that is selected im guessing

print(wrkbk.sheetnames)	#prints the workbooks sheet names

#get the names of the headings
# for i in range(1, sh.max_column+1): print(sh.cell(row=1, column=i).value)



# find a way to get the headings and their values into seperate places on the word document

# print(sh.cell(row=1, column=1).value)	#get one of the values that are assigned to the thingy


# iterate through excel and display data, goes row by row but may be more benificial to go column for column
for i in range(1, sh.max_row+1):
   # print("\n")
   
   # doc.add_heading(f"Headings",level=1)	#no need for headings any more but in case you ever do need them the way to get them is here	
   if i != 1:
   	   doc.add_heading(f"Employee {i - 1} data:",level=1)	#-1 as we used one for the headings and the rest would have one added on

   	   for j in range(1, sh.max_column+1):
   	   		cell_obj = sh.cell(row=i, column=j)
   	   		doc.add_paragraph(f"{sh.cell(row=1, column=j).value} : { sh.cell(row=i, column=j).value}",style='Normal')
       # print(cell_obj.value, end=" ")

wrkbk.close()

#gets the structure of the text & prints it 
# for paragraph in doc.paragraphs:
	# for run in paragraph.runs:
		# print("\n",run.text)	#print all the paragraphs from the word file
    
doc.save('./InputChecker.docx')#save the appended word document